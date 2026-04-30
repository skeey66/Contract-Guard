"""분석 결과를 기반으로 수정안이 반영된 계약서를 DOCX/PDF/HWPX 형식으로 생성.

위험 조항(high/medium)에 수정안이 있으면 그것을, 없으면 원문을 그대로 사용해
바로 사용 가능한 깨끗한 계약서 형태로 출력한다 (위험도 라벨·원문 비교 표기 없음).
"""

import io
import logging
import os
import zipfile
from datetime import datetime
from pathlib import Path
from xml.sax.saxutils import escape as xml_escape

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from backend.app.models.analysis import AnalysisResult, ClauseAnalysis
from backend.app.models.risk import RiskLevel

logger = logging.getLogger(__name__)

_KOREAN_FONT_CANDIDATES = [
    # macOS
    "/System/Library/Fonts/Supplemental/AppleGothic.ttf",
    "/Library/Fonts/AppleGothic.ttf",
    "/System/Library/Fonts/AppleSDGothicNeo.ttc",
    # Linux
    "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
    # Windows
    "C:/Windows/Fonts/malgun.ttf",
    "C:/Windows/Fonts/malgunbd.ttf",
]
_PDF_FONT_NAME = "KoreanFont"
_PDF_FONT_REGISTERED = False


def _ensure_korean_font_for_pdf() -> bool:
    """reportlab에 한국어 폰트를 1회 등록. 성공 시 True."""
    global _PDF_FONT_REGISTERED
    if _PDF_FONT_REGISTERED:
        return True
    for path in _KOREAN_FONT_CANDIDATES:
        if os.path.exists(path):
            try:
                if path.endswith(".ttc"):
                    pdfmetrics.registerFont(TTFont(_PDF_FONT_NAME, path, subfontIndex=0))
                else:
                    pdfmetrics.registerFont(TTFont(_PDF_FONT_NAME, path))
                _PDF_FONT_REGISTERED = True
                logger.info(f"PDF 한국어 폰트 등록: {path}")
                return True
            except Exception as e:
                logger.warning(f"폰트 등록 실패 ({path}): {e}")
                continue
    logger.error("한국어 폰트를 찾을 수 없습니다 (PDF 한글이 깨질 수 있음)")
    return False


def _is_risky(ca: ClauseAnalysis) -> bool:
    return ca.risk_level in (RiskLevel.HIGH, RiskLevel.MEDIUM)


def _final_rewrite(ca: ClauseAnalysis) -> tuple[str | None, str]:
    """수정안 본문과 라벨을 우선순위에 따라 반환.

    우선순위: user_override(사용자 직접 수정) > suggested_rewrite(LLM 권고안)
    수정안이 없으면 (None, "")을 반환한다.
    """
    if ca.user_override and ca.user_override.strip():
        return ca.user_override, "[수정안 (사용자)]"
    if ca.suggested_rewrite and ca.suggested_rewrite.strip():
        return ca.suggested_rewrite, "[권고 수정안]"
    return None, ""


def _final_clause_text(ca: ClauseAnalysis) -> str:
    """최종 계약서에 들어갈 조항 본문을 반환.

    위험 조항(high/medium)에 수정안이 있으면 수정안을, 없으면 원문을 그대로 사용한다.
    안전·저위험 조항은 항상 원문 유지.
    """
    if _is_risky(ca):
        rewrite_text, _ = _final_rewrite(ca)
        if rewrite_text:
            return rewrite_text
    return ca.clause_content or ""


def _risk_label(level: RiskLevel) -> str:
    return {
        RiskLevel.HIGH: "고위험",
        RiskLevel.MEDIUM: "중위험",
        RiskLevel.LOW: "저위험",
        RiskLevel.SAFE: "안전",
    }.get(level, str(level))


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def build_docx(result: AnalysisResult) -> bytes:
    """수정안이 반영된 계약서 DOCX를 생성하여 바이트로 반환."""
    doc = Document()

    # 기본 폰트 (맑은 고딕)
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)

    title = doc.add_heading("권고 수정안 반영 계약서", level=1)
    for run in title.runs:
        run.font.name = "맑은 고딕"

    meta = doc.add_paragraph()
    meta.add_run(f"원본 파일: {result.filename}\n").italic = True
    meta.add_run(f"생성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n").italic = True
    meta.add_run(f"전체 조항 {result.total_clauses}개 / 위험 조항 {result.risky_clauses}개").italic = True

    doc.add_paragraph(
        "본 문서는 AI 분석 결과를 바탕으로, 위험 요소가 식별된 조항에 권고 수정안을 반영하여 "
        "재구성한 참고용 계약서입니다. 법적 효력은 없으며, 실제 계약 체결 전 전문가 검토를 권장합니다."
    )
    doc.add_paragraph()

    for ca in result.clause_analyses:
        # 조항 제목
        h = doc.add_heading(f"제{ca.clause_index}조 {ca.clause_title}".strip(), level=2)
        for run in h.runs:
            run.font.name = "맑은 고딕"

        # 본문 — 위험 조항이면 수정안, 아니면 원문
        final_text = _final_clause_text(ca)
        for line in final_text.splitlines() or [final_text]:
            doc.add_paragraph(line)

        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF (reportlab)
# ---------------------------------------------------------------------------

def _pdf_styles() -> dict:
    """PDF용 스타일 사전. 한국어 폰트가 없으면 기본 폰트로 폴백."""
    has_korean = _ensure_korean_font_for_pdf()
    base_font = _PDF_FONT_NAME if has_korean else "Helvetica"

    title = ParagraphStyle(
        "TitleK", fontName=base_font, fontSize=18, leading=24, spaceAfter=14
    )
    meta = ParagraphStyle(
        "MetaK", fontName=base_font, fontSize=9, leading=14, textColor="#666666"
    )
    h2 = ParagraphStyle(
        "H2K", fontName=base_font, fontSize=13, leading=18, spaceBefore=14, spaceAfter=6, textColor="#222222"
    )
    body = ParagraphStyle(
        "BodyK", fontName=base_font, fontSize=10.5, leading=16
    )
    risk_high = ParagraphStyle("RiskHigh", parent=body, textColor="#C0392B")
    risk_med = ParagraphStyle("RiskMed", parent=body, textColor="#D38B1E")
    risk_low = ParagraphStyle("RiskLow", parent=body, textColor="#888888")
    risk_safe = ParagraphStyle("RiskSafe", parent=body, textColor="#2E7D32")
    rewrite = ParagraphStyle(
        "RewriteK", parent=body, textColor="#1F4E8B", leftIndent=8
    )
    label = ParagraphStyle(
        "LabelK", fontName=base_font, fontSize=10.5, leading=16, textColor="#555555"
    )
    return {
        "title": title,
        "meta": meta,
        "h2": h2,
        "body": body,
        "high": risk_high,
        "medium": risk_med,
        "low": risk_low,
        "safe": risk_safe,
        "rewrite": rewrite,
        "label": label,
    }


def _para(text: str, style: ParagraphStyle) -> Paragraph:
    """줄바꿈을 <br/>로 치환한 reportlab Paragraph."""
    safe = xml_escape(text or "").replace("\n", "<br/>")
    return Paragraph(safe, style)


def build_pdf(result: AnalysisResult) -> bytes:
    """수정안이 반영된 계약서 PDF를 생성하여 바이트로 반환."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
    )
    styles = _pdf_styles()
    story = []

    story.append(_para("권고 수정안 반영 계약서", styles["title"]))
    meta_text = (
        f"원본 파일: {result.filename}<br/>"
        f"생성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}<br/>"
        f"전체 조항 {result.total_clauses}개 / 위험 조항 {result.risky_clauses}개"
    )
    story.append(Paragraph(meta_text, styles["meta"]))
    story.append(Spacer(1, 0.4 * cm))
    story.append(_para(
        "본 문서는 AI 분석 결과를 바탕으로, 위험 요소가 식별된 조항에 권고 수정안을 반영하여 "
        "재구성한 참고용 계약서입니다. 법적 효력은 없으며, 실제 계약 체결 전 전문가 검토를 권장합니다.",
        styles["body"],
    ))
    story.append(Spacer(1, 0.6 * cm))

    for ca in result.clause_analyses:
        story.append(_para(f"제{ca.clause_index}조 {ca.clause_title}".strip(), styles["h2"]))
        final_text = _final_clause_text(ca)
        story.append(_para(final_text, styles["body"]))
        story.append(Spacer(1, 0.4 * cm))

    doc.build(story)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# HWPX (최소 OWPML 구조의 ZIP 생성)
# ---------------------------------------------------------------------------
# HWPX는 OWPML 기반 ZIP 컨테이너. 한컴오피스가 열 수 있는 최소한의 파일 셋만 생성한다.
# 정식 스펙은 매우 복잡하므로, 본문 단락만 포함한 간이 형태로 작성한다.


_HWPX_MIMETYPE = "application/hwp+zip"

_HWPX_VERSION_XML = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
    '<hv:HCFVersion xmlns:hv="http://www.hancom.co.kr/hwpml/2011/version" '
    'tagetApplication="HWPML" major="5" minor="1" micro="0" buildNumber="0" os="32"/>\n'
)

_HWPX_CONTAINER_XML = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
    '<ocf:container xmlns:ocf="urn:oasis:names:tc:opendocument:xmlns:container" version="1.0">\n'
    '  <ocf:rootfiles>\n'
    '    <ocf:rootfile full-path="Contents/content.hpf" media-type="application/hwpml-package+xml"/>\n'
    '  </ocf:rootfiles>\n'
    '</ocf:container>\n'
)

_HWPX_HPF = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
    '<opf:package xmlns:opf="http://www.idpf.org/2007/opf/" version="1.2" unique-identifier="hwpx-uid">\n'
    '  <opf:metadata xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:opf-met="http://www.idpf.org/2007/opf/meta">\n'
    '    <opf-met:meta name="title">권고 수정안 반영 계약서</opf-met:meta>\n'
    '  </opf:metadata>\n'
    '  <opf:manifest>\n'
    '    <opf:item id="header" href="header.xml" media-type="application/xml"/>\n'
    '    <opf:item id="section0" href="section0.xml" media-type="application/xml"/>\n'
    '  </opf:manifest>\n'
    '  <opf:spine>\n'
    '    <opf:itemref idref="section0" linear="yes"/>\n'
    '  </opf:spine>\n'
    '</opf:package>\n'
)

_HWPX_HEADER = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
    '<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" version="1.2" secCnt="1">\n'
    '</hh:head>\n'
)


def _hwpx_paragraph(text: str, *, bold: bool = False, color: str | None = None) -> str:
    """HWPX <hp:p> 단락을 생성. 색상은 RRGGBB 16진수."""
    safe = xml_escape(text or "")
    bold_attr = ' bold="1"' if bold else ""
    color_attr = f' textColor="#{color}"' if color else ""
    # 매우 단순화된 단락. 한글 오피스가 미지원 속성은 무시한다.
    return (
        '  <hp:p id="0" paraPrIDRef="0" styleIDRef="0">\n'
        f'    <hp:run charPrIDRef="0"{bold_attr}{color_attr}><hp:t>{safe}</hp:t></hp:run>\n'
        '  </hp:p>\n'
    )


def _build_hwpx_section(result: AnalysisResult) -> str:
    """본문 section0.xml을 생성."""
    body = []
    body.append(_hwpx_paragraph("권고 수정안 반영 계약서", bold=True))
    body.append(_hwpx_paragraph(f"원본 파일: {result.filename}"))
    body.append(_hwpx_paragraph(f"생성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}"))
    body.append(_hwpx_paragraph(
        f"전체 조항 {result.total_clauses}개 / 위험 조항 {result.risky_clauses}개"
    ))
    body.append(_hwpx_paragraph(""))
    body.append(_hwpx_paragraph(
        "본 문서는 AI 분석 결과를 바탕으로, 위험 요소가 식별된 조항에 권고 수정안을 반영하여 "
        "재구성한 참고용 계약서입니다. 법적 효력은 없으며, 실제 계약 체결 전 전문가 검토를 권장합니다."
    ))
    body.append(_hwpx_paragraph(""))

    for ca in result.clause_analyses:
        body.append(_hwpx_paragraph(f"제{ca.clause_index}조 {ca.clause_title}".strip(), bold=True))
        final_text = _final_clause_text(ca)
        for line in final_text.splitlines() or [final_text]:
            body.append(_hwpx_paragraph(line))
        body.append(_hwpx_paragraph(""))

    section_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" '
        'xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">\n'
        + "".join(body) +
        '</hs:sec>\n'
    )
    return section_xml


def build_hwpx(result: AnalysisResult) -> bytes:
    """HWPX 컨테이너를 ZIP으로 패키징하여 반환.

    한컴오피스가 열 수 있는 최소 OWPML 구조만 포함한다.
    복잡한 스타일·표·그림 등은 지원하지 않는다.
    """
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        # mimetype은 비압축으로 가장 먼저 들어가야 한다 (OCF 규약)
        info = zipfile.ZipInfo("mimetype")
        info.compress_type = zipfile.ZIP_STORED
        zf.writestr(info, _HWPX_MIMETYPE)

        zf.writestr("version.xml", _HWPX_VERSION_XML)
        zf.writestr("META-INF/container.xml", _HWPX_CONTAINER_XML)
        zf.writestr("Contents/content.hpf", _HWPX_HPF)
        zf.writestr("Contents/header.xml", _HWPX_HEADER)
        zf.writestr("Contents/section0.xml", _build_hwpx_section(result))
    return buf.getvalue()


# ---------------------------------------------------------------------------
# 통합 export 진입점
# ---------------------------------------------------------------------------

_EXPORT_HANDLERS = {
    "docx": (build_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    "pdf": (build_pdf, "application/pdf"),
    "hwpx": (build_hwpx, "application/hwp+zip"),
}


def export_analysis(result: AnalysisResult, fmt: str) -> tuple[bytes, str, str]:
    """주어진 형식으로 계약서를 생성하여 (바이트, MIME, 파일명)을 반환."""
    fmt_norm = (fmt or "").lower().strip()
    handler = _EXPORT_HANDLERS.get(fmt_norm)
    if not handler:
        raise ValueError(f"지원하지 않는 형식: {fmt}")
    builder, mime = handler
    data = builder(result)
    base = Path(result.filename).stem or "contract"
    out_name = f"{base}_권고수정안.{fmt_norm}"
    return data, mime, out_name


def supported_formats() -> list[str]:
    return list(_EXPORT_HANDLERS.keys())
