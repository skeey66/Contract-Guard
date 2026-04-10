from pathlib import Path
import fitz  # PyMuPDF
from docx import Document
from hwp2yaml import extract_hwp_text

# 지원 확장자 → 추출 함수 매핑
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".hwp", ".hwpx"}


def extract_text(file_path: str) -> tuple[str, int]:
  """파일 확장자에 따라 텍스트를 추출한다. (전체 텍스트, 페이지/단락 수)를 반환."""
  ext = Path(file_path).suffix.lower()
  if ext == ".pdf":
    return _extract_pdf(file_path)
  if ext == ".docx":
    return _extract_docx(file_path)
  if ext in (".hwp", ".hwpx"):
    return _extract_hwp(file_path)
  raise ValueError(f"지원하지 않는 파일 형식: {ext}")


def _extract_pdf(path: str) -> tuple[str, int]:
  doc = fitz.open(path)
  pages = [page.get_text() for page in doc]
  doc.close()
  return "\n".join(pages), len(pages)


def _extract_docx(path: str) -> tuple[str, int]:
  doc = Document(path)
  blocks: list[str] = []

  # 본문 단락
  for p in doc.paragraphs:
    text = p.text.strip()
    if text:
      blocks.append(text)

  # 표 내부 텍스트 (계약서 조항이 표로 작성된 경우 대응)
  for table in doc.tables:
    for row in table.rows:
      for cell in row.cells:
        for p in cell.paragraphs:
          text = p.text.strip()
          if text:
            blocks.append(text)

  return "\n".join(blocks), len(blocks)


def _extract_hwp(path: str) -> tuple[str, int]:
  result = extract_hwp_text(path)
  if not result.success:
    raise ValueError(f"HWP 텍스트 추출 실패: {result.error}")

  lines = [line.strip() for line in result.text.split("\n") if line.strip()]
  return "\n".join(lines), len(lines)
